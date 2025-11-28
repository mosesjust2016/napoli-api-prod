from flask_openapi3 import APIBlueprint, Tag
from flask import jsonify, request, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity, get_jwt
from sqlalchemy import func, desc, and_, or_
from datetime import datetime, timedelta
from pydantic import BaseModel, Field, EmailStr, field_validator, model_validator
from typing import List, Optional, Dict, Any, Union
import os
import secrets
import base64
import uuid
from docx import Document
from num2words import num2words
from werkzeug.utils import secure_filename

# Fixed imports
from core.addons.extensions import db
from core.models.employees import Employee
from core.models.companies import Company
from core.models.hr_actions import HRAction
from core.models.employee_documents import EmployeeDocument
from core.models.users import User
from core.models.auditLogModel import AuditLog

# Google Drive imports
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
import pickle

employee_tag = Tag(name="Employees", description="Employee management operations")
employee_bp = APIBlueprint(
    'employee', __name__, url_prefix='/api/employees', abp_tags=[employee_tag]
)

# ---------------------- GOOGLE DRIVE CONFIGURATION ---------------------- #
SCOPES = ['https://www.googleapis.com/auth/drive']
# Root folder ID for Napoli HR
PARENT_FOLDER_ID = "1pN7V0ngbP8EMcz1FXf2QhXEiFH1gbccF"

def get_drive_service():
    """Get authenticated Google Drive service"""
    try:
        creds = None
        # Token file should be in the project root or config directory
        token_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'token.json')
        
        if os.path.exists(token_path):
            creds = Credentials.from_authorized_user_file(token_path, SCOPES)
        
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                creds.refresh(Request())
            else:
                # For server environments, you might need service account credentials
                # This would need to be configured based on your deployment
                raise Exception("Google Drive credentials not configured properly")
            
            # Save the credentials for the next run
            with open(token_path, 'w') as token:
                token.write(creds.to_json())
        
        return build('drive', 'v3', credentials=creds)
    except Exception as e:
        print(f"Error initializing Google Drive service: {str(e)}")
        # Fallback to local storage if Drive fails
        return None

def create_drive_folder(service, folder_name, parent_id=None):
    """Create a folder in Google Drive"""
    try:
        if not service:
            raise Exception("Google Drive service not available")
            
        file_metadata = {
            'name': folder_name,
            'mimeType': 'application/vnd.google-apps.folder'
        }
        if parent_id:
            file_metadata['parents'] = [parent_id]
        
        folder = service.files().create(body=file_metadata, fields='id,name,webViewLink').execute()
        print(f"DEBUG: Created Google Drive folder: {folder_name} with ID: {folder.get('id')}")
        return folder
    except Exception as e:
        print(f"Error creating Google Drive folder: {str(e)}")
        raise

def upload_to_drive(service, file_path, file_name, parent_id=None):
    """Upload a file to Google Drive"""
    try:
        if not service:
            raise Exception("Google Drive service not available")
            
        file_metadata = {
            'name': file_name
        }
        if parent_id:
            file_metadata['parents'] = [parent_id]
        
        media = MediaFileUpload(file_path, resumable=True)
        file = service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id,name,webViewLink,webContentLink'
        ).execute()
        
        print(f"DEBUG: Uploaded file to Google Drive: {file_name} with ID: {file.get('id')}")
        return file
    except Exception as e:
        print(f"Error uploading file to Google Drive: {str(e)}")
        raise

def find_or_create_employee_folder(service, employee, company):
    """Find or create employee folder structure in Google Drive"""
    try:
        # First, ensure company folder exists
        company_folder_name = f"{company.name.replace(' ', '_')}_{company.id}"
        company_folder = None
        
        # Search for existing company folder
        query = f"name='{company_folder_name}' and mimeType='application/vnd.google-apps.folder' and '{PARENT_FOLDER_ID}' in parents and trashed=false"
        results = service.files().list(q=query, fields="files(id, name)").execute()
        folders = results.get('files', [])
        
        if folders:
            company_folder = folders[0]
        else:
            company_folder = create_drive_folder(service, company_folder_name, PARENT_FOLDER_ID)
        
        # Create employee folder inside company folder
        employee_folder_name = f"{employee.first_name}_{employee.last_name}_{employee.employee_id}".replace(' ', '_')
        employee_folder_query = f"name='{employee_folder_name}' and mimeType='application/vnd.google-apps.folder' and '{company_folder['id']}' in parents and trashed=false"
        employee_results = service.files().list(q=employee_folder_query, fields="files(id, name)").execute()
        employee_folders = employee_results.get('files', [])
        
        if employee_folders:
            employee_folder = employee_folders[0]
        else:
            employee_folder = create_drive_folder(service, employee_folder_name, company_folder['id'])
            
            # Create subfolders
            subfolders = ['Personal_Documents', 'Employment_Documents', 'Certificates', 'HR_Actions']
            for subfolder in subfolders:
                create_drive_folder(service, subfolder, employee_folder['id'])
        
        return employee_folder
        
    except Exception as e:
        print(f"Error creating employee folder structure in Google Drive: {str(e)}")
        raise

# ---------------------- EMPLOYEE ID GENERATION ---------------------- #
def generate_employee_id(company_id: int) -> str:
    """
    Generate employee ID in format: {COMPANY_PREFIX}{SEQUENCE_NUMBER}
    Example: AMA001, AMA002, NAP001, etc.
    """
    try:
        # Get company prefix from database
        company = Company.query.get(company_id)
        if not company:
            raise ValueError(f"Company with ID {company_id} not found")
        
        # Get company prefix - use custom prefix if available, otherwise generate from name
        if hasattr(company, 'employee_id_prefix') and company.employee_id_prefix:
            prefix = company.employee_id_prefix.upper()
        else:
            # Generate prefix from company name (first 3 characters uppercase)
            prefix = company.name[:3].upper().replace(' ', '')
            # Ensure we have at least 3 characters
            if len(prefix) < 3:
                prefix = prefix.ljust(3, 'X')
        
        # Get the last employee ID for this company
        last_employee = Employee.query.filter(
            Employee.company_id == company_id,
            Employee.employee_id.like(f"{prefix}%")
        ).order_by(Employee.id.desc()).first()
        
        if last_employee and last_employee.employee_id:
            # Extract the numeric part and increment
            last_id = last_employee.employee_id
            # Remove prefix and get numeric part
            numeric_part = last_id.replace(prefix, "")
            try:
                next_number = int(numeric_part) + 1
            except ValueError:
                next_number = 1
        else:
            # First employee for this company
            next_number = 1
        
        # Format with leading zeros (AMA001, AMA002, etc.)
        return f"{prefix}{next_number:03d}"
        
    except Exception as e:
        print(f"Error generating employee ID: {str(e)}")
        # Fallback: use timestamp-based ID
        return f"EMP{int(datetime.now().timestamp())}"

# ---------------------- UPDATED DOCUMENT HANDLING FUNCTIONS ---------------------- #
def create_documents_directory(employee_id: int) -> str:
    """Create local directory structure for employee documents (temporary storage)"""
    try:
        # Base documents directory - using Napoli HR Folders structure
        base_dir = "/app/Napoli HR Folders"
        os.makedirs(base_dir, exist_ok=True)
        
        # Employee-specific directory
        employee_dir = os.path.join(base_dir, f"Employee_{employee_id}")
        os.makedirs(employee_dir, exist_ok=True)
            
        return employee_dir
        
    except Exception as e:
        print(f"Error creating documents directory: {str(e)}")
        return "/app/Napoli HR Folders"

def handle_employee_documents(employee_id: int, documents_data: Dict[str, Any], uploaded_by: int, employee, company) -> List[EmployeeDocument]:
    """Handle saving uploaded documents to Google Drive and database"""
    saved_documents = []
    
    try:
        # Initialize Google Drive service
        drive_service = get_drive_service()
        if not drive_service:
            print("Google Drive service not available, falling back to local storage")
            return handle_employee_documents_local(employee_id, documents_data, uploaded_by)
        
        # Find or create employee folder in Google Drive
        employee_folder = find_or_create_employee_folder(drive_service, employee, company)
        
        # Create temporary local directory
        employee_dir = create_documents_directory(employee_id)
        
        # Document type mapping and folder mapping
        document_type_mapping = {
            'profilePhoto': 'id_card',
            'nrcCopy': 'id_card',  
            'cv': 'resume',
            'offerLetter': 'contract',
            'certificates': 'certificate'
        }
        
        document_folder_mapping = {
            'id_card': 'Personal_Documents',
            'resume': 'Employment_Documents', 
            'contract': 'Employment_Documents',
            'certificate': 'Certificates'
        }
        
        for doc_key, doc_data in documents_data.items():
            try:
                # Handle single documents
                if doc_key in ['profilePhoto', 'nrcCopy', 'cv', 'offerLetter'] and doc_data:
                    saved_doc = save_base64_document_drive(
                        employee_id=employee_id,
                        base64_data=doc_data,
                        document_type=document_type_mapping[doc_key],
                        document_name=f"{document_type_mapping[doc_key].replace('_', ' ').title()}",
                        uploaded_by=uploaded_by,
                        employee_dir=employee_dir,
                        doc_key=doc_key,
                        drive_service=drive_service,
                        employee_folder=employee_folder,
                        folder_mapping=document_folder_mapping
                    )
                    if saved_doc:
                        saved_documents.append(saved_doc)
                
                # Handle certificates array
                elif doc_key == 'certificates' and doc_data:
                    for i, cert_data in enumerate(doc_data):
                        if cert_data:
                            saved_cert = save_base64_document_drive(
                                employee_id=employee_id,
                                base64_data=cert_data,
                                document_type='certificate',
                                document_name=f"Certificate {i+1}",
                                uploaded_by=uploaded_by,
                                employee_dir=employee_dir,
                                doc_key=f"certificate_{i+1}",
                                drive_service=drive_service,
                                employee_folder=employee_folder,
                                folder_mapping=document_folder_mapping
                            )
                            if saved_cert:
                                saved_documents.append(saved_cert)
                                
            except Exception as e:
                print(f"Error processing document {doc_key}: {str(e)}")
                continue
        
        # Clean up temporary local files
        try:
            import shutil
            shutil.rmtree(employee_dir)
            print(f"DEBUG: Cleaned up temporary directory: {employee_dir}")
        except Exception as cleanup_error:
            print(f"Warning: Could not clean up temporary directory: {str(cleanup_error)}")
                
        return saved_documents
        
    except Exception as e:
        print(f"Error in Google Drive document handling, falling back to local: {str(e)}")
        return handle_employee_documents_local(employee_id, documents_data, uploaded_by)

def handle_employee_documents_local(employee_id: int, documents_data: Dict[str, Any], uploaded_by: int) -> List[EmployeeDocument]:
    """Fallback function to handle documents locally"""
    saved_documents = []
    employee_dir = create_documents_directory(employee_id)
    
    document_type_mapping = {
        'profilePhoto': 'id_card',
        'nrcCopy': 'id_card',  
        'cv': 'resume',
        'offerLetter': 'contract',
        'certificates': 'certificate'
    }
    
    for doc_key, doc_data in documents_data.items():
        try:
            if doc_key in ['profilePhoto', 'nrcCopy', 'cv', 'offerLetter'] and doc_data:
                saved_doc = save_base64_document(
                    employee_id=employee_id,
                    base64_data=doc_data,
                    document_type=document_type_mapping[doc_key],
                    document_name=f"{document_type_mapping[doc_key].replace('_', ' ').title()}",
                    uploaded_by=uploaded_by,
                    employee_dir=employee_dir,
                    doc_key=doc_key
                )
                if saved_doc:
                    saved_documents.append(saved_doc)
            
            elif doc_key == 'certificates' and doc_data:
                for i, cert_data in enumerate(doc_data):
                    if cert_data:
                        saved_cert = save_base64_document(
                            employee_id=employee_id,
                            base64_data=cert_data,
                            document_type='certificate',
                            document_name=f"Certificate {i+1}",
                            uploaded_by=uploaded_by,
                            employee_dir=employee_dir,
                            doc_key=f"certificate_{i+1}"
                        )
                        if saved_cert:
                            saved_documents.append(saved_cert)
                            
        except Exception as e:
            print(f"Error processing document {doc_key} locally: {str(e)}")
            continue
    
    return saved_documents

def save_base64_document_drive(employee_id: int, base64_data: str, document_type: str, 
                              document_name: str, uploaded_by: int, employee_dir: str, 
                              doc_key: str, drive_service, employee_folder, folder_mapping) -> Optional[EmployeeDocument]:
    """Save base64 document to Google Drive and database"""
    try:
        # Extract file extension from base64 data
        if base64_data.startswith('data:'):
            header = base64_data.split(';')[0]
            mime_type = header.split(':')[1] if ':' in header else header
            
            mime_to_extension = {
                'image/png': 'png',
                'image/jpeg': 'jpg',
                'image/jpg': 'jpg',
                'application/pdf': 'pdf',
                'application/msword': 'doc',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx'
            }
            
            file_extension = mime_to_extension.get(mime_type, 'bin')
            base64_data = base64_data.split(',')[1]
        else:
            file_extension = 'bin'
        
        # Fix base64 padding issues
        padding = len(base64_data) % 4
        if padding:
            base64_data += '=' * (4 - padding)
        
        # Decode base64 data
        file_data = base64.b64decode(base64_data)
        
        # Generate filename and save locally temporarily
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{doc_key}_{employee_id}_{timestamp}.{file_extension}"
        local_file_path = os.path.join(employee_dir, filename)
        
        # Save file locally temporarily
        with open(local_file_path, 'wb') as f:
            f.write(file_data)
        
        # Upload to Google Drive
        target_folder_name = folder_mapping.get(document_type, 'Personal_Documents')
        
        # Find the target subfolder
        query = f"name='{target_folder_name}' and mimeType='application/vnd.google-apps.folder' and '{employee_folder['id']}' in parents and trashed=false"
        results = drive_service.files().list(q=query, fields="files(id, name)").execute()
        target_folders = results.get('files', [])
        
        if target_folders:
            target_folder_id = target_folders[0]['id']
        else:
            # Create subfolder if it doesn't exist
            target_folder = create_drive_folder(drive_service, target_folder_name, employee_folder['id'])
            target_folder_id = target_folder['id']
        
        # Upload file to Google Drive
        drive_file = upload_to_drive(drive_service, local_file_path, filename, target_folder_id)
        
        # Save to database with Google Drive URL
        document = EmployeeDocument(
            employee_id=employee_id,
            document_type=document_type,
            document_name=document_name,
            file_url=drive_file.get('webViewLink'),  # Store Google Drive view link
            file_drive_id=drive_file.get('id'),  # Store Google Drive file ID
            upload_date=datetime.now(),
            uploaded_by=uploaded_by,
            is_verified=True,
            verified_by=uploaded_by,
            comments=f"Uploaded to Google Drive during employee creation: {document_name}",
            created_at=datetime.now(),
            updated_at=datetime.now()
        )
        
        db.session.add(document)
        db.session.flush()
        
        print(f"DEBUG: Saved document {document_name} to Google Drive: {drive_file.get('webViewLink')}")
        
        # Clean up local file
        try:
            os.remove(local_file_path)
        except Exception as cleanup_error:
            print(f"Warning: Could not clean up local file {local_file_path}: {str(cleanup_error)}")
        
        return document
        
    except Exception as e:
        print(f"Error saving base64 document to Google Drive {document_name}: {str(e)}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        return None

def save_base64_document(employee_id: int, base64_data: str, document_type: str, 
                        document_name: str, uploaded_by: int, employee_dir: str, doc_key: str) -> Optional[EmployeeDocument]:
    """Original function for local storage (fallback)"""
    try:
        if base64_data.startswith('data:'):
            header = base64_data.split(';')[0]
            mime_type = header.split(':')[1] if ':' in header else header
            
            mime_to_extension = {
                'image/png': 'png',
                'image/jpeg': 'jpg',
                'image/jpg': 'jpg',
                'application/pdf': 'pdf',
                'application/msword': 'doc',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx'
            }
            
            file_extension = mime_to_extension.get(mime_type, 'bin')
            base64_data = base64_data.split(',')[1]
        else:
            file_extension = 'bin'
        
        padding = len(base64_data) % 4
        if padding:
            base64_data += '=' * (4 - padding)
        
        file_data = base64.b64decode(base64_data)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{doc_key}_{employee_id}_{timestamp}.{file_extension}"
        file_path = os.path.join(employee_dir, filename)
        
        with open(file_path, 'wb') as f:
            f.write(file_data)
        
        document = EmployeeDocument(
            employee_id=employee_id,
            document_type=document_type,
            document_name=document_name,
            file_url=file_path,
            upload_date=datetime.now(),
            uploaded_by=uploaded_by,
            is_verified=True,
            verified_by=uploaded_by,
            comments=f"Uploaded during employee creation: {document_name}",
            created_at=datetime.now(),
            updated_at=datetime.now()
        )
        
        db.session.add(document)
        db.session.flush()
        
        print(f"DEBUG: Saved document {document_name} locally to {file_path}")
        return document
        
    except Exception as e:
        print(f"Error saving base64 document locally {document_name}: {str(e)}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        return None

def generate_documents_from_templates(employee: Employee, company: Company, documents_data: Dict[str, Any] = None) -> Dict[str, str]:
    """Generate documents using template files and save to Google Drive"""
    documents_data = documents_data or {}
    
    # Create temporary local directory
    employee_dir = create_documents_directory(employee.id)
    generated_docs = {
        'new_joiner_form': None,
        'employment_contract': None,
        'documents_folder': employee_dir
    }
    
    try:
        # Initialize Google Drive service
        drive_service = get_drive_service()
        if drive_service:
            employee_folder = find_or_create_employee_folder(drive_service, employee, company)
        else:
            employee_folder = None
            print("Google Drive service not available, saving documents locally")
    except Exception as e:
        print(f"Google Drive initialization failed, using local storage: {str(e)}")
        drive_service = None
        employee_folder = None
    
    # Updated template paths to be relative to the project root
    base_dir = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
    new_joiner_template_path = os.path.join(base_dir, "templates", "NEW JOINER FORM 2025.pdf")
    contract_template_path = os.path.join(base_dir, "templates", "Contract Template.pdf")
    
    # Generate New Joiner Form from template
    try:
        if os.path.exists(new_joiner_template_path):
            new_joiner_doc = Document(new_joiner_template_path)
            
            # Prepare replacements for new joiner form
            new_joiner_replacements = {
                'Employee Name:': f'Employee Name: {employee.first_name} {employee.middle_name or ""} {employee.last_name}'.strip(),
                'Date of birth\n(dd/mm/yyyy):': f'Date of birth\n(dd/mm/yyyy): {employee.date_of_birth.strftime("%d/%m/%Y") if employee.date_of_birth else ""}',
                'NRC Number:': f'NRC Number: {employee.national_id or ""}',
                'Employing Company:': f'Employing Company: {company.name}',
                'Hire Date:': f'Hire Date: {employee.start_date.strftime("%d/%m/%Y") if employee.start_date else ""}',
                'Job Title:': f'Job Title: {employee.position or ""}',
                'NAPSA:': f'NAPSA: {employee.pension_number or ""}',
                'NHIMA:': f'NHIMA: {documents_data.get("nhima_number", "")}',
                'TPIN:': f'TPIN: {employee.tax_id or ""}',
                'Bank Name:': f'Bank Name: {employee.bank_name or ""}',
                'Branch:': f'Branch: {documents_data.get("bank_branch", "")}',
                'Account\nNumber:': f'Account\nNumber: {employee.bank_account or ""}',
                'Sort\nCode:': f'Sort\nCode: {documents_data.get("sort_code", "")}',
                'Gender (M/F):': f'Gender (M/F): {employee.gender[0] if employee.gender else ""}',
                'Phone:': f'Phone: {employee.phone or ""}',
                'Home Address:': f'Home Address: {employee.address or ""}',
                'Marital\nStatus\n(Single /\nMarried /\nDivorced):': f'Marital\nStatus\n(Single /\nMarried /\nDivorced): {employee.marital_status or ""}',
                'Spouse\nName:': f'Spouse\nName: {documents_data.get("spouse_name", "")}',
                'Next of Kin:': f'Next of Kin: {employee.emergency_contact_name or ""}',
                'Children\nNames & DOB:': f'Children\nNames & DOB: {documents_data.get("children", "")}'
            }
            
            for paragraph in new_joiner_doc.paragraphs:
                for old_text, new_text in new_joiner_replacements.items():
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, str(new_text))
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"New_Joiner_Form_{employee.id}_{timestamp}.docx"
            local_filepath = os.path.join(employee_dir, filename)
            new_joiner_doc.save(local_filepath)
            generated_docs['new_joiner_form'] = local_filepath
            
            # Upload to Google Drive if available
            if drive_service and employee_folder:
                try:
                    # Find Employment_Documents folder
                    query = f"name='Employment_Documents' and mimeType='application/vnd.google-apps.folder' and '{employee_folder['id']}' in parents and trashed=false"
                    results = drive_service.files().list(q=query, fields="files(id, name)").execute()
                    target_folders = results.get('files', [])
                    
                    if target_folders:
                        target_folder_id = target_folders[0]['id']
                    else:
                        target_folder = create_drive_folder(drive_service, 'Employment_Documents', employee_folder['id'])
                        target_folder_id = target_folder['id']
                    
                    drive_file = upload_to_drive(drive_service, local_filepath, filename, target_folder_id)
                    generated_docs['new_joiner_form_drive'] = drive_file.get('webViewLink')
                    generated_docs['new_joiner_form_drive_id'] = drive_file.get('id')
                    
                    # Update the file path to use Drive URL for database storage
                    generated_docs['new_joiner_form'] = drive_file.get('webViewLink')
                    
                except Exception as drive_error:
                    print(f"Failed to upload New Joiner Form to Google Drive: {str(drive_error)}")
            
            print(f"DEBUG: New Joiner Form saved to: {local_filepath}")
        else:
            print(f"Warning: New Joiner template not found at {new_joiner_template_path}")
        
    except Exception as e:
        print(f"Error generating New Joiner Form: {str(e)}")
    
    # Generate Employment Contract from template
    try:
        if os.path.exists(contract_template_path):
            contract_doc = Document(contract_template_path)
            
            # Calculate salary components
            basic_salary = float(employee.salary) if employee.salary else 0
            housing_allowance = documents_data.get("housing_allowance", 0) or 0
            transport_allowance = documents_data.get("transport_allowance", 0) or 0
            lunch_allowance = documents_data.get("lunch_allowance", 0) or 0
            fuel_allowance = documents_data.get("fuel_allowance", 0) or 0
            phone_allowance = documents_data.get("phone_allowance", 0) or 0
            
            total_salary = basic_salary + housing_allowance + transport_allowance + lunch_allowance + fuel_allowance + phone_allowance
            amount_words = num2words(total_salary, lang="en").title() if total_salary > 0 else "Zero"
            
            # Prepare replacements for contract
            contract_replacements = {
                '[Company]': company.name,
                '[Employee]': f'{employee.first_name} {employee.last_name}',
                '[X]': datetime.now().strftime('%d'),
                '[month year]': datetime.now().strftime('%B %Y'),
                '[Job Title]': employee.position or '[Job Title]',
                '[XX date]': employee.start_date.strftime('%Y-%m-%d') if employee.start_date else '[Start Date]',
                '[Amount]': f'{total_salary:.2f}' if total_salary > 0 else '[Amount]',
                '[Write Amount Out in Full]': f'{amount_words} Kwacha',
                'Basic Salary ZMW [Amount]': f'Basic Salary ZMW {basic_salary:.2f}' if basic_salary > 0 else 'Basic Salary ZMW [Amount]',
                'Housing Allowance ZMW [Amount]': f'Housing Allowance ZMW {housing_allowance:.2f}' if housing_allowance > 0 else 'Housing Allowance ZMW [Amount]',
                'Transport Allowance ZMW [Amount]': f'Transport Allowance ZMW {transport_allowance:.2f}' if transport_allowance > 0 else 'Transport Allowance ZMW [Amount]',
                'Lunch Allowance ZMW [Amount]': f'Lunch Allowance ZMW {lunch_allowance:.2f}' if lunch_allowance > 0 else 'Lunch Allowance ZMW [Amount]',
                '[Fuel Allowance]': f'Fuel Allowance ZMW {fuel_allowance:.2f}' if fuel_allowance > 0 else '',
                '[Phone Allowance]': f'Phone Allowance ZMW {phone_allowance:.2f}' if phone_allowance > 0 else '',
                '[Company Vehicle]': documents_data.get("company_vehicle", '') if documents_data.get("company_vehicle") else '',
                '[Company Phone]': documents_data.get("company_phone", '') if documents_data.get("company_phone") else '',
                '[Company-Provided Accommodation (include address)]': documents_data.get("company_accommodation", '') if documents_data.get("company_accommodation") else '',
                '[Monday to Friday, from 08:00 to 17:00 and Saturday, from 08:00 to 13:00.]': documents_data.get("working_hours", "Monday to Friday, from 08:00 to 17:00 and Saturday, from 08:00 to 13:00."),
                '«Currency»': employee.salary_currency,
                '«Lunch_Allowance_Figure»': str(lunch_allowance) if lunch_allowance > 0 else '[Amount]',
                '[add the below manually, as not standard offering]:': '',
                '[[OR]]': '',
                '[OR]': ''
            }
            
            for paragraph in contract_doc.paragraphs:
                for old_text, new_text in contract_replacements.items():
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, str(new_text))
            
            if not fuel_allowance:
                for paragraph in contract_doc.paragraphs:
                    if 'Fuel Allowance' in paragraph.text and not fuel_allowance:
                        paragraph.clear()
            
            if not phone_allowance:
                for paragraph in contract_doc.paragraphs:
                    if 'Phone Allowance' in paragraph.text and not phone_allowance:
                        paragraph.clear()
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"Employment_Contract_{employee.id}_{timestamp}.docx"
            local_filepath = os.path.join(employee_dir, filename)
            contract_doc.save(local_filepath)
            generated_docs['employment_contract'] = local_filepath
            
            # Upload to Google Drive if available
            if drive_service and employee_folder:
                try:
                    # Find Employment_Documents folder
                    query = f"name='Employment_Documents' and mimeType='application/vnd.google-apps.folder' and '{employee_folder['id']}' in parents and trashed=false"
                    results = drive_service.files().list(q=query, fields="files(id, name)").execute()
                    target_folders = results.get('files', [])
                    
                    if target_folders:
                        target_folder_id = target_folders[0]['id']
                    else:
                        target_folder = create_drive_folder(drive_service, 'Employment_Documents', employee_folder['id'])
                        target_folder_id = target_folder['id']
                    
                    drive_file = upload_to_drive(drive_service, local_filepath, filename, target_folder_id)
                    generated_docs['employment_contract_drive'] = drive_file.get('webViewLink')
                    generated_docs['employment_contract_drive_id'] = drive_file.get('id')
                    
                    # Update the file path to use Drive URL for database storage
                    generated_docs['employment_contract'] = drive_file.get('webViewLink')
                    
                except Exception as drive_error:
                    print(f"Failed to upload Employment Contract to Google Drive: {str(drive_error)}")
            
            print(f"DEBUG: Employment Contract saved to: {local_filepath}")
        else:
            print(f"Warning: Contract template not found at {contract_template_path}")
        
    except Exception as e:
        print(f"Error generating Employment Contract: {str(e)}")
    
    # Clean up temporary local directory
    try:
        import shutil
        shutil.rmtree(employee_dir)
        print(f"DEBUG: Cleaned up temporary directory: {employee_dir}")
    except Exception as cleanup_error:
        print(f"Warning: Could not clean up temporary directory: {str(cleanup_error)}")
    
    return generated_docs

def save_document_to_database(employee_id: int, document_type: str, file_path: str, uploaded_by: int, document_name: str, comments: str = None, expiry_date: datetime = None, drive_file_id: str = None) -> Optional[EmployeeDocument]:
    """Save document record to database with Google Drive support"""
    try:
        document_type_mapping = {
            'new_joiner_form': 'other',
            'employment_contract': 'contract'
        }
        
        db_document_type = document_type_mapping.get(document_type, 'other')
        
        document = EmployeeDocument(
            employee_id=employee_id,
            document_type=db_document_type,
            document_name=document_name,
            file_url=file_path,
            file_drive_id=drive_file_id,
            upload_date=datetime.now(),
            uploaded_by=uploaded_by,
            expiry_date=expiry_date,
            is_verified=True,
            verified_by=uploaded_by,
            comments=comments or f"Automatically generated {document_name}",
            created_at=datetime.now(),
            updated_at=datetime.now()
        )
        
        db.session.add(document)
        db.session.flush()
        print(f"DEBUG: Document saved to database: {document_name} with type: {db_document_type}")
        return document
        
    except Exception as e:
        print(f"Error saving document to database: {str(e)}")
        db.session.rollback()
        return None

# ---------------------- SCHEMAS ---------------------- #
class EmployeeResponseSchema(BaseModel):
    employee_id: str = Field(..., description="Employee ID (e.g., AMA001)")
    id: int = Field(..., description="Employee ID")
    first_name: str = Field(..., description="First name")
    last_name: str = Field(..., description="Last name")
    email: str = Field(..., description="Email")
    phone: str = Field(..., description="Phone")
    personal_email: Optional[str] = Field(None, description="Personal email")
    date_of_birth: str = Field(..., description="Date of birth")
    national_id: Optional[str] = Field(None, description="National ID")
    work_permit_number: Optional[str] = Field(None, description="Work Permit Number")
    identity_type: str = Field(..., description="Identity Type (NRC or Work Permit)")
    work_permit_valid_from: Optional[str] = Field(None, description="Work permit valid from date")
    work_permit_valid_to: Optional[str] = Field(None, description="Work permit valid to date")
    work_permit_expiry_notified: bool = Field(False, description="Whether work permit expiry notification was sent")
    identity_document_is_expired: bool = Field(False, description="Whether work permit is expired")
    identity_document_expires_soon: bool = Field(False, description="Whether work permit expires soon")
    primary_identity_number: Optional[str] = Field(None, description="Primary identity number based on type")
    requires_work_permit_renewal: bool = Field(False, description="Whether work permit requires renewal")
    gender: str = Field(..., description="Gender")
    marital_status: Optional[str] = Field(None, description="Marital status")
    address: Optional[str] = Field(None, description="Address")
    emergency_contact_name: str = Field(..., description="Emergency contact name")
    emergency_contact_phone: str = Field(..., description="Emergency contact phone")
    emergency_contact_relationship: str = Field(..., description="Emergency contact relationship")
    company_id: int = Field(..., description="Company ID")
    department: str = Field(..., description="Department")
    position: str = Field(..., description="Position")
    employment_type: str = Field(..., description="Employment type")
    employment_status: str = Field(..., description="Employment status")
    start_date: str = Field(..., description="Start date")
    end_date: Optional[str] = Field(None, description="End date")
    probation_end_date: Optional[str] = Field(None, description="Probation end date")
    contract_end_date: Optional[str] = Field(None, description="Contract end date")
    supervisor_id: Optional[int] = Field(None, description="Supervisor ID")
    work_location: Optional[str] = Field(None, description="Work location")
    salary: float = Field(..., description="Salary")
    salary_currency: str = Field(..., description="Salary currency")
    payment_frequency: str = Field(..., description="Payment frequency")
    bank_name: Optional[str] = Field(None, description="Bank name")
    bank_account: Optional[str] = Field(None, description="Bank account")
    tax_id: Optional[str] = Field(None, description="Tax ID")
    pension_number: Optional[str] = Field(None, description="Pension number")
    has_live_disciplinary: bool = Field(..., description="Has live disciplinary")
    created_by: int = Field(..., description="Created by")
    created_at: str = Field(..., description="Creation timestamp")
    updated_at: str = Field(..., description="Update timestamp")
    company_name: Optional[str] = Field(None, description="Company name")
    nationality: Optional[str] = Field(None, description="Nationality")


class EmployeeCreateSchema(BaseModel):
    first_name: str = Field(..., min_length=1, description="First name")
    last_name: str = Field(..., min_length=1, description="Last name")
    date_of_birth: str = Field(..., description="Date of birth")
    
    # Identity Document Fields
    identity_type: str = Field("NRC", description="Identity Type (NRC or Work Permit)")
    national_id: Optional[str] = Field(None, min_length=1, description="National ID - required if identity_type is NRC")
    work_permit_number: Optional[str] = Field(None, description="Work Permit Number - required if identity_type is Work Permit")
    work_permit_valid_from: Optional[str] = Field(None, description="Work permit valid from date - required if identity_type is Work Permit")
    work_permit_valid_to: Optional[str] = Field(None, description="Work permit valid to date - required if identity_type is Work Permit")
    
    # Nationality field for conditional validation
    nationality: str = Field("Zambia", description="Nationality of the employee")
    
    gender: str = Field(..., description="Gender")
    phone: Optional[str] = Field(None, description="Phone")
    personal_email: Optional[EmailStr] = Field(None, description="Personal email")
    address: Optional[str] = Field(None, description="Address")
    marital_status: Optional[str] = Field(None, description="Marital status")
    company_id: int = Field(..., description="Company ID")
    position: str = Field(..., min_length=1, description="Position")
    department: str = Field(..., description="Department")
    employment_type: str = Field(..., description="Employment type")
    employment_status: str = Field("Active", description="Employment status")
    start_date: str = Field(..., description="Start date")
    end_date: Optional[str] = Field(None, description="End date")
    probation_end_date: Optional[str] = Field(None, description="Probation end date")
    contract_end_date: Optional[str] = Field(None, description="Contract end date")
    supervisor_id: Optional[int] = Field(None, description="Supervisor ID")
    work_location: Optional[str] = Field(None, description="Work location")
    salary: Optional[float] = Field(None, ge=0, description="Salary")
    salary_currency: str = Field("ZMW", description="Salary currency")
    payment_frequency: str = Field("Monthly", description="Payment frequency")
    bank_name: Optional[str] = Field(None, description="Bank name")
    bank_account: Optional[str] = Field(None, description="Bank account")
    tax_id: Optional[str] = Field(None, description="Tax ID")
    pension_number: Optional[str] = Field(None, description="Pension number")
    generate_documents: Optional[bool] = Field(True, description="Generate onboarding documents")

    # Emergency contact fields
    emergency_contact_name: Optional[str] = Field("Not Provided", description="Emergency contact name")
    emergency_contact_phone: Optional[str] = Field("Not Provided", description="Emergency contact phone")
    emergency_contact_relationship: Optional[str] = Field("Not Provided", description="Emergency contact relationship")
    
    # Email is now optional and will be generated if not provided
    email: Optional[EmailStr] = Field(None, description="Email")

    # Additional fields for document generation
    middle_name: Optional[str] = Field(None, description="Middle name")
    napsa_number: Optional[str] = Field(None, description="NAPSA number")
    nhima_number: Optional[str] = Field(None, description="NHIMA number")
    tpin: Optional[str] = Field(None, description="TPIN number")
    account_number: Optional[str] = Field(None, description="Bank account number")
    sort_code: Optional[str] = Field(None, description="Bank sort code")
    next_of_kin: Optional[str] = Field(None, description="Next of kin")
    physical_address: Optional[str] = Field(None, description="Physical address")

    # Document generation fields
    housing_allowance: Optional[float] = Field(0, description="Housing allowance")
    transport_allowance: Optional[float] = Field(0, description="Transport allowance")
    lunch_allowance: Optional[float] = Field(0, description="Lunch allowance")
    fuel_allowance: Optional[float] = Field(0, description="Fuel allowance")
    phone_allowance: Optional[float] = Field(0, description="Phone allowance")
    company_vehicle: Optional[str] = Field(None, description="Company vehicle details")
    company_phone: Optional[str] = Field(None, description="Company phone details")
    company_accommodation: Optional[str] = Field(None, description="Company accommodation details")
    working_hours: Optional[str] = Field("Monday to Friday, from 08:00 to 17:00 and Saturday, from 08:00 to 13:00.", description="Working hours")
    bank_branch: Optional[str] = Field(None, description="Bank branch")
    spouse_name: Optional[str] = Field(None, description="Spouse name")
    children: Optional[str] = Field(None, description="Children details")

    # Documents field
    documents: Optional[Dict[str, Any]] = Field(default_factory=dict, description="Employee documents")

    # Validators to convert integers to strings for phone and account fields
    @field_validator('phone', 'emergency_contact_phone', 'bank_account', 'account_number', 'sort_code', mode='before')
    @classmethod
    def convert_to_string(cls, v):
        if v is None:
            return v
        return str(v)

    @field_validator('email', 'personal_email')
    @classmethod
    def email_to_lowercase(cls, v):
        if v:
            return v.strip().lower()
        return v

    @field_validator('identity_type')
    @classmethod
    def validate_identity_type(cls, v):
        valid_types = ['NRC', 'Work Permit']
        if v not in valid_types:
            raise ValueError(f'Identity type must be one of: {", ".join(valid_types)}')
        return v

    @field_validator('nationality')
    @classmethod
    def validate_nationality(cls, v):
        if not v:
            return "Zambia"
        return v

    # FIXED: Employment type validator to match database ENUM
    @field_validator('employment_type')
    @classmethod
    def validate_employment_type(cls, v):
        if v is None:
            return v
                
        valid_types = [
            'FULL-TIME', 'PART-TIME', 'CONTRACT', 'PERMANENT', 
            'Full-time', 'Part-time', 'Contract', 'Permanent',
            'FIXED-TERM', 'Fixed Term', 'Fixed-Term', 'FIXED TERM',
            'INTERN', 'Intern', 'APPRENTICE', 'Apprentice',
            'CONSULTANT', 'Consultant',
            'PROBATION', 'Probation'
        ]

        if v not in valid_types:
            raise ValueError(f'Employment type must be one of: {", ".join(valid_types)}')

        # Normalize to match the database ENUM exactly
        v_lower = v.lower()
        if v_lower in ['permanent', 'full-time', 'fulltime']:
            return 'Full-time'
        elif v_lower in ['part-time', 'parttime']:
            return 'Part-time'
        elif v_lower in ['contract']:
            return 'Contract'
        elif v_lower in ['fixed-term', 'fixed term', 'fixed_term']:
            return 'Fixed-Term'
        elif v_lower in ['intern', 'internship']:
            return 'Intern'
        elif v_lower in ['apprentice', 'apprenticeship']:
            return 'Apprentice'
        elif v_lower in ['consultant', 'consultancy']:
            return 'Consultant'
        elif v_lower in ['probation']:
            return 'Probation'
        return v

    @model_validator(mode='after')
    def validate_all_fields(self):
        """Comprehensive validation for all fields"""
        # Normalize nationality check
        is_zambian = self.nationality.lower() in ['zambia', 'zambian']
        
        # Validate identity type matches nationality
        if is_zambian:
            if self.identity_type != 'NRC':
                raise ValueError('Zambian employees must use NRC as identity type')
            if not self.national_id:
                raise ValueError('National ID is required for Zambian employees')
            # Clear work permit fields for Zambians
            self.work_permit_number = None
            self.work_permit_valid_from = None
            self.work_permit_valid_to = None
        else:
            if self.identity_type != 'Work Permit':
                raise ValueError('Non-Zambian employees must use Work Permit as identity type')
            if not self.work_permit_number:
                raise ValueError('Work Permit Number is required for non-Zambian employees')
            # Clear national_id for non-Zambians
            self.national_id = None

        # Set default email if not provided
        if not self.email:
            email_username = f"{self.first_name.lower().replace(' ', '.')}.{self.last_name.lower()}"
            self.email = f"{email_username}@company.com"

        # Set default department if not provided
        if not self.department:
            self.department = "General"

        # Ensure documents is never None
        if self.documents is None:
            self.documents = {}

        return self

    class Config:
        extra = 'ignore'

class EmployeeUpdateSchema(BaseModel):
    first_name: Optional[str] = Field(None, min_length=1, description="First name")
    last_name: Optional[str] = Field(None, min_length=1, description="Last name")
    email: Optional[EmailStr] = Field(None, description="Email")
    phone: Optional[Union[str, int]] = Field(None, description="Phone")
    personal_email: Optional[EmailStr] = Field(None, description="Personal email")
    date_of_birth: Optional[str] = Field(None, description="Date of birth")
    national_id: Optional[str] = Field(None, min_length=1, description="National ID")
    work_permit_number: Optional[str] = Field(None, description="Work Permit Number")
    identity_type: Optional[str] = Field(None, description="Identity Type (NRC or Work Permit)")
    work_permit_valid_from: Optional[str] = Field(None, description="Work permit valid from date")
    work_permit_valid_to: Optional[str] = Field(None, description="Work permit valid to date")
    work_permit_expiry_notified: Optional[bool] = Field(None, description="Whether work permit expiry notification was sent")
    nationality: Optional[str] = Field(None, description="Nationality")
    gender: Optional[str] = Field(None, description="Gender")
    marital_status: Optional[str] = Field(None, description="Marital status")
    address: Optional[str] = Field(None, description="Address")
    emergency_contact_name: Optional[str] = Field(None, description="Emergency contact name")
    emergency_contact_phone: Optional[Union[str, int]] = Field(None, description="Emergency contact phone")
    emergency_contact_relationship: Optional[str] = Field(None, description="Emergency contact relationship")
    company_id: Optional[int] = Field(None, description="Company ID")
    department: Optional[str] = Field(None, description="Department")
    position: Optional[str] = Field(None, min_length=1, description="Position")
    employment_type: Optional[str] = Field(None, description="Employment type")
    employment_status: Optional[str] = Field(None, description="Employment status")
    start_date: Optional[str] = Field(None, description="Start date")
    end_date: Optional[str] = Field(None, description="End date")
    probation_end_date: Optional[str] = Field(None, description="Probation end date")
    contract_end_date: Optional[str] = Field(None, description="Contract end date")
    supervisor_id: Optional[int] = Field(None, description="Supervisor ID")
    work_location: Optional[str] = Field(None, description="Work location")
    salary: Optional[float] = Field(None, ge=0, description="Salary")
    salary_currency: Optional[str] = Field(None, description="Salary currency")
    payment_frequency: Optional[str] = Field(None, description="Payment frequency")
    bank_name: Optional[str] = Field(None, description="Bank name")
    bank_account: Optional[Union[str, int]] = Field(None, description="Bank account")
    tax_id: Optional[str] = Field(None, description="Tax ID")
    pension_number: Optional[str] = Field(None, description="Pension number")
    force_update: Optional[bool] = Field(False, description="Force update")

    # Validators to convert integers to strings for phone and account fields
    @field_validator('phone', 'emergency_contact_phone', 'bank_account', mode='before')
    @classmethod
    def convert_to_string(cls, v):
        if v is None:
            return v
        return str(v)

    # FIXED: Employment type validator to match database ENUM
    @field_validator('employment_type')
    @classmethod
    def validate_employment_type(cls, v):
        if v is None:
            return v
                
        valid_types = [
            'FULL-TIME', 'PART-TIME', 'CONTRACT', 'PERMANENT', 
            'Full-time', 'Part-time', 'Contract', 'Permanent',
            'FIXED-TERM', 'Fixed Term', 'Fixed-Term', 'FIXED TERM',
            'INTERN', 'Intern', 'APPRENTICE', 'Apprentice',
            'CONSULTANT', 'Consultant',
            'PROBATION', 'Probation'
        ]

        if v not in valid_types:
            raise ValueError(f'Employment type must be one of: {", ".join(valid_types)}')

        # Normalize to match the database ENUM exactly
        v_lower = v.lower()
        if v_lower in ['permanent', 'full-time', 'fulltime']:
            return 'Full-time'
        elif v_lower in ['part-time', 'parttime']:
            return 'Part-time'
        elif v_lower in ['contract']:
            return 'Contract'
        elif v_lower in ['fixed-term', 'fixed term', 'fixed_term']:
            return 'Fixed-Term'
        elif v_lower in ['intern', 'internship']:
            return 'Intern'
        elif v_lower in ['apprentice', 'apprenticeship']:
            return 'Apprentice'
        elif v_lower in ['consultant', 'consultancy']:
            return 'Consultant'
        elif v_lower in ['probation']:
            return 'Probation'
        return v

class DocumentGenerationResponse(BaseModel):
    employee_id: int = Field(..., description="Employee ID")
    new_joiner_form_url: Optional[str] = Field(None, description="New Joiner Form URL")
    contract_url: Optional[str] = Field(None, description="Employment Contract URL")
    message: str = Field(..., description="Response message")

class EmployeeListResponseSchema(BaseModel):
    employees: List[EmployeeResponseSchema] = Field(..., description="List of employees")
    pagination: Dict[str, Any] = Field(..., description="Pagination information")

class EmployeeFullResponseSchema(BaseModel):
    employee: EmployeeResponseSchema = Field(..., description="Employee information")
    hr_actions: List[Dict[str, Any]] = Field(..., description="HR actions")
    disciplinary_records: List[Dict[str, Any]] = Field(..., description="Disciplinary records")
    leave_records: List[Dict[str, Any]] = Field(..., description="Leave records")
    documents: List[Dict[str, Any]] = Field(..., description="Documents")

class WorkPermitExpiryResponse(BaseModel):
    expired_work_permits: List[Dict[str, Any]] = Field(..., description="List of employees with expired work permits")
    expiring_soon_work_permits: List[Dict[str, Any]] = Field(..., description="List of employees with work permits expiring soon")
    expired_count: int = Field(..., description="Count of expired work permits")
    expiring_soon_count: int = Field(..., description="Count of expiring soon work permits")
    check_date: str = Field(..., description="Date when the check was performed")

class ErrorResponse(BaseModel):
    status: int = Field(..., description="HTTP status code")
    isError: bool = Field(True, description="Indicates if the response is an error")
    message: str = Field(..., description="Error message")

class SuccessResponse(BaseModel):
    message: str = Field(..., description="Success message")
    status: int = Field(200, description="HTTP status code")
    isError: bool = Field(False, description="Indicates if the response is an error")

# ---------------------- BULK UPLOAD SCHEMAS ---------------------- #
class BulkEmployeeCreateSchema(BaseModel):
    employees: List[EmployeeCreateSchema] = Field(..., description="List of employees to create")
    skip_errors: bool = Field(False, description="Skip employees with errors and continue processing")
    send_notifications: bool = Field(False, description="Send email notifications to new employees")

class BulkEmployeeResponseSchema(BaseModel):
    total_processed: int = Field(..., description="Total employees processed")
    successful: int = Field(..., description="Number of successfully created employees")
    failed: int = Field(..., description="Number of failed creations")
    errors: List[Dict[str, Any]] = Field(..., description="List of errors for failed creations")
    successful_employees: List[Dict[str, Any]] = Field(..., description="List of successfully created employees")

# Path parameter schemas
class EmployeeIdPath(BaseModel):
    employee_id: int = Field(..., description="Employee ID")

# ---------------------- DATE PARSING HELPER ---------------------- #
def parse_date(date_str):
    """Parse date from string or Excel serial number"""
    if not date_str:
        return None
    
    # Handle Excel serial numbers
    try:
        excel_serial = float(date_str)
        if 0 <= excel_serial <= 100000:
            base_date = datetime(1899, 12, 30)
            parsed_date = base_date + timedelta(days=excel_serial)
            return parsed_date.date()
    except (ValueError, TypeError):
        pass
    
    # Handle standard date formats
    date_formats = [
        '%Y-%m-%d', '%m/%d/%Y', '%d/%m/%Y', '%d-%m-%Y', '%m-%d-%Y', '%Y/%m/%d'
    ]
    
    for date_format in date_formats:
        try:
            return datetime.strptime(date_str, date_format).date()
        except ValueError:
            continue
    
    raise ValueError(f"Invalid date format: {date_str}")

# ---------------------- ROUTES ---------------------- #
@employee_bp.get('/', responses={"200": EmployeeListResponseSchema, "500": ErrorResponse}, security=[{"jwt": []}])
@jwt_required()
def get_employees():
    """Get paginated list of employees with filtering and sorting"""
    try:
        # Pagination parameters
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        per_page = min(per_page, 100)

        # Base query
        query = Employee.query

        # Filtering
        status = request.args.get('status', 'all')
        company_id = request.args.get('company_id')
        search = request.args.get('search', '')
        identity_type = request.args.get('identity_type', 'all')
        work_permit_status = request.args.get('work_permit_status', 'all')
        nationality = request.args.get('nationality', 'all')

        # Status filter
        if status == 'active':
            query = query.filter(Employee.employment_status.in_(['Active', 'Probation']))
        elif status in ['Active', 'Probation', 'Inactive', 'Expired Contract']:
            query = query.filter_by(employment_status=status)

        # Company filter
        if company_id and company_id != 'all':
            query = query.filter_by(company_id=company_id)

        # Identity type filter
        if identity_type and identity_type != 'all':
            query = query.filter_by(identity_type=identity_type)

        # Nationality filter
        if nationality and nationality != 'all':
            query = query.filter_by(nationality=nationality)

        # Work permit status filter
        if work_permit_status != 'all':
            today = datetime.now().date()
            if work_permit_status == 'expired':
                query = query.filter(
                    Employee.identity_type == 'Work Permit',
                    Employee.work_permit_valid_to < today
                )
            elif work_permit_status == 'expiring_soon':
                thirty_days_from_now = today + timedelta(days=30)
                query = query.filter(
                    Employee.identity_type == 'Work Permit',
                    Employee.work_permit_valid_to >= today,
                    Employee.work_permit_valid_to <= thirty_days_from_now
                )
            elif work_permit_status == 'valid':
                query = query.filter(
                    Employee.identity_type == 'Work Permit',
                    Employee.work_permit_valid_to > today + timedelta(days=30)
                )

        # Search filter
        if search:
            search_term = f"%{search}%"
            query = query.filter(
                or_(
                    Employee.first_name.ilike(search_term),
                    Employee.last_name.ilike(search_term),
                    Employee.email.ilike(search_term),
                    Employee.phone.ilike(search_term),
                    Employee.national_id.ilike(search_term),
                    Employee.work_permit_number.ilike(search_term),
                    Employee.position.ilike(search_term),
                    Employee.employee_id.ilike(search_term),
                    Employee.nationality.ilike(search_term)
                )
            )

        # Sorting
        sort_by = request.args.get('sort_by', 'first_name')
        sort_order = request.args.get('sort_order', 'asc')
        
        if sort_order == 'desc':
            query = query.order_by(desc(sort_by))
        else:
            query = query.order_by(sort_by)

        # Execute paginated query
        pagination = query.paginate(
            page=page, 
            per_page=per_page, 
            error_out=False
        )

        employees = pagination.items

        # Add disciplinary flag for display
        employees_data = []
        for employee in employees:
            emp_data = employee.to_dict()
            emp_data['has_live_disciplinary_flag'] = '🔴' if employee.has_live_disciplinary else ''
            employees_data.append(emp_data)

        return jsonify({
            "employees": employees_data,
            "pagination": {
                "page": page,
                "per_page": per_page,
                "total": pagination.total,
                "pages": pagination.pages,
                "has_next": pagination.has_next,
                "has_prev": pagination.has_prev
            }
        }), 200

    except Exception as e:
        return jsonify({
            "status": 500,
            "isError": True,
            "message": f"Error fetching employees: {str(e)}"
        }), 500

@employee_bp.post(
    '/create',
    responses={"201": EmployeeResponseSchema, "400": ErrorResponse, "403": ErrorResponse, "409": ErrorResponse, "500": ErrorResponse},
    security=[{"jwt": []}]
)
@jwt_required()
def create_employee(body: EmployeeCreateSchema):
    """Create a new employee with Google Drive document storage"""
    try:
        current_user_id = get_jwt_identity()
        
        # Check HR admin access
        jwt_data = get_jwt()
        roles = jwt_data.get('roles', [])
        
        if 'hr_admin' not in roles and 'admin' not in roles:
            return jsonify({
                "status": 403,
                "isError": True,
                "message": "HR admin or Super Admin access required"
            }), 403

        # Validate company exists
        company = Company.query.get(body.company_id)
        if not company:
            return jsonify({
                "status": 404,
                "isError": True,
                "message": "Company not found"
            }), 404

        # Validate supervisor exists if provided
        if body.supervisor_id:
            supervisor = Employee.query.get(body.supervisor_id)
            if not supervisor:
                return jsonify({
                    "status": 400,
                    "isError": True,
                    "message": f"Supervisor with ID {body.supervisor_id} not found. Please create the supervisor first or set supervisor_id to null."
                }), 400

        # Generate employee ID
        employee_id = generate_employee_id(body.company_id)

        # Check if generated employee ID already exists
        existing_employee = Employee.query.filter_by(employee_id=employee_id).first()
        if existing_employee:
            employee_id = f"EMP{int(datetime.now().timestamp())}"

        # Check if email already exists
        if body.email:
            existing_employee = Employee.query.filter_by(email=body.email).first()
            if existing_employee:
                return jsonify({
                    "status": 409,
                    "isError": True,
                    "message": "An employee with this email already exists"
                }), 409
        
        # ========== CRITICAL FIX: VALIDATE NATIONALITY AND IDENTITY TYPE MATCH ==========
        is_zambian = body.nationality and body.nationality.lower() in ['zambia', 'zambian']
        
        if is_zambian and body.identity_type != 'NRC':
            return jsonify({
                "status": 400,
                "isError": True,
                "message": "Zambian employees must use NRC as identity type"
            }), 400
        
        if not is_zambian and body.identity_type != 'Work Permit':
            return jsonify({
                "status": 400,
                "isError": True,
                "message": "Non-Zambian employees must use Work Permit as identity type"
            }), 400
        
        # Check identity document uniqueness based on nationality and identity type
        if is_zambian:
            if body.national_id:
                existing_national_id = Employee.query.filter_by(national_id=body.national_id).first()
                if existing_national_id:
                    return jsonify({
                        "status": 409,
                        "isError": True,
                        "message": "An employee with this National ID already exists"
                    }), 409
            if not body.national_id:
                return jsonify({
                    "status": 400,
                    "isError": True,
                    "message": "National ID is required for Zambian employees"
                }), 400
        else:
            if body.work_permit_number:
                existing_work_permit = Employee.query.filter_by(work_permit_number=body.work_permit_number).first()
                if existing_work_permit:
                    return jsonify({
                        "status": 409,
                        "isError": True,
                        "message": "An employee with this Work Permit Number already exists"
                    }), 409
            if not body.work_permit_number:
                return jsonify({
                    "status": 400,
                    "isError": True,
                    "message": "Work Permit Number is required for non-Zambian employees"
                }), 400
            if not body.work_permit_valid_from or not body.work_permit_valid_to:
                return jsonify({
                    "status": 400,
                    "isError": True,
                    "message": "Work Permit valid from and valid to dates are required for non-Zambian employees"
                }), 400

        # Parse dates
        date_of_birth = parse_date(body.date_of_birth)
        start_date = parse_date(body.start_date)
        end_date = parse_date(body.end_date) if body.end_date else None
        probation_end_date = parse_date(body.probation_end_date) if body.probation_end_date else None
        contract_end_date = parse_date(body.contract_end_date) if body.contract_end_date else None
        
        work_permit_valid_from = None
        work_permit_valid_to = None
        if not is_zambian and body.identity_type == 'Work Permit':
            work_permit_valid_from = parse_date(body.work_permit_valid_from) if body.work_permit_valid_from else None
            work_permit_valid_to = parse_date(body.work_permit_valid_to) if body.work_permit_valid_to else None

        salary = body.salary if body.salary is not None else 0.0

        # Create employee
        employee = Employee(
            employee_id=employee_id,
            first_name=body.first_name,
            last_name=body.last_name,
            email=body.email,
            phone=body.phone or "Not Provided",
            personal_email=body.personal_email,
            date_of_birth=date_of_birth,
            
            # Identity fields
            national_id=body.national_id,
            work_permit_number=body.work_permit_number,
            identity_type=body.identity_type,
            work_permit_valid_from=work_permit_valid_from,
            work_permit_valid_to=work_permit_valid_to,
            work_permit_expiry_notified=False,
            nationality=body.nationality,
            
            # Additional fields
            middle_name=body.middle_name,
            napsa_number=body.napsa_number,
            nhima_number=body.nhima_number,
            tpin=body.tax_id,
            account_number=body.bank_account,
            sort_code=body.sort_code,
            next_of_kin=body.next_of_kin,
            physical_address=body.physical_address,
            
            gender=body.gender,
            marital_status=body.marital_status,
            address=body.address,
            emergency_contact_name=body.emergency_contact_name or "Not Provided",
            emergency_contact_phone=body.emergency_contact_phone or "Not Provided",
            emergency_contact_relationship=body.emergency_contact_relationship or "Not Provided",
            company_id=body.company_id,
            department=body.department,
            position=body.position,
            employment_type=body.employment_type,
            employment_status=body.employment_status,
            start_date=start_date,
            end_date=end_date,
            probation_end_date=probation_end_date,
            contract_end_date=contract_end_date,
            supervisor_id=body.supervisor_id,
            work_location=body.work_location,
            salary=salary,
            salary_currency=body.salary_currency,
            payment_frequency=body.payment_frequency,
            bank_name=body.bank_name,
            bank_account=body.bank_account,
            tax_id=body.tax_id,
            pension_number=body.pension_number,
            created_by=int(current_user_id)
        )
        
        db.session.add(employee)
        db.session.flush()
        
        # ========== UPDATED: HANDLE UPLOADED DOCUMENTS WITH GOOGLE DRIVE ==========
        if body.documents:
            try:
                saved_documents = handle_employee_documents(
                    employee_id=employee.id,
                    documents_data=body.documents,
                    uploaded_by=current_user_id,
                    employee=employee,
                    company=company
                )
                print(f"DEBUG: Saved {len(saved_documents)} documents for employee {employee.id}")
            except Exception as doc_error:
                print(f"Warning: Document processing failed but employee created: {str(doc_error)}")
                import traceback
                print(f"Document error traceback: {traceback.format_exc()}")
        # ========== END UPDATED DOCUMENTS HANDLING ==========
        
        # Generate template documents if requested
        if body.generate_documents:
            try:
                documents_data = {
                    "nhima_number": body.nhima_number,
                    "bank_branch": body.bank_branch,
                    "sort_code": body.sort_code,
                    "spouse_name": body.spouse_name,
                    "children": body.children,
                    "housing_allowance": body.housing_allowance or 0,
                    "transport_allowance": body.transport_allowance or 0,
                    "lunch_allowance": body.lunch_allowance or 0,
                    "fuel_allowance": body.fuel_allowance or 0,
                    "phone_allowance": body.phone_allowance or 0,
                    "company_vehicle": body.company_vehicle,
                    "company_phone": body.company_phone,
                    "company_accommodation": body.company_accommodation,
                    "working_hours": body.working_hours
                }
                
                generated_docs = generate_documents_from_templates(employee, company, documents_data)
                
                # Save documents to database with Google Drive support
                if generated_docs.get('new_joiner_form'):
                    drive_file_id = generated_docs.get('new_joiner_form_drive_id')
                    save_document_to_database(
                        employee_id=employee.id,
                        document_type='new_joiner_form',
                        file_path=generated_docs['new_joiner_form'],
                        uploaded_by=current_user_id,
                        document_name=f"New Joiner Form - {employee.first_name} {employee.last_name}",
                        comments="Employee onboarding form with personal and employment details",
                        drive_file_id=drive_file_id
                    )
                
                if generated_docs.get('employment_contract'):
                    expiry_date = None
                    if employee.contract_end_date:
                        expiry_date = employee.contract_end_date
                    elif employee.start_date:
                        expiry_date = employee.start_date + timedelta(days=365)
                    
                    drive_file_id = generated_docs.get('employment_contract_drive_id')
                    save_document_to_database(
                        employee_id=employee.id,
                        document_type='employment_contract',
                        file_path=generated_docs['employment_contract'],
                        uploaded_by=current_user_id,
                        document_name=f"Employment Contract - {employee.first_name} {employee.last_name}",
                        comments="Formal employment agreement with terms and conditions",
                        expiry_date=expiry_date,
                        drive_file_id=drive_file_id
                    )
                    
            except Exception as doc_error:
                print(f"Template document generation failed but employee created: {str(doc_error)}")
        
        # Update company employee count
        company.employee_count = Employee.query.filter_by(
            company_id=body.company_id
        ).filter(
            Employee.employment_status.in_(['Active', 'Probation'])
        ).count()
        
        db.session.commit()
        
        # Create audit log
        audit = AuditLog(
            employee_id=employee.id,
            action="CREATE",
            performed_by=current_user_id,
            details=f"Employee {employee.employee_id} created successfully with nationality: {employee.nationality}, identity type: {employee.identity_type}. Documents stored in Google Drive."
        )
        db.session.add(audit)
        db.session.commit()
        
        employee_data = employee.to_dict()
        if 'company_name' not in employee_data or not employee_data['company_name']:
            employee_data['company_name'] = company.name
            
        return jsonify(employee_data), 201

    except Exception as e:
        db.session.rollback()
        print(f"Error in create_employee: {str(e)}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        return jsonify({
            "status": 500,
            "isError": True,
            "message": f"Error creating employee: {str(e)}"
        }), 500

@employee_bp.get('/documents/<path:filename>')
@jwt_required()
def serve_document(filename):
    """Serve generated documents from Napoli HR Folders"""
    try:
        documents_base = "/app/Napoli HR Folders"
        file_path = os.path.join(documents_base, filename)
        
        # If filename doesn't include full path, try to find it in employee folders
        if not os.path.exists(file_path):
            # Look for file in any employee folder
            for folder in os.listdir(documents_base):
                folder_path = os.path.join(documents_base, folder)
                if os.path.isdir(folder_path) and folder.startswith('Employee_'):
                    potential_path = os.path.join(folder_path, filename)
                    if os.path.exists(potential_path):
                        file_path = potential_path
                        break
        
        if os.path.exists(file_path) and os.path.isfile(file_path):
            return send_file(file_path, as_attachment=False)
        else:
            return jsonify({
                "status": 404,
                "isError": True,
                "message": "Document not found"
            }), 404
            
    except Exception as e:
        return jsonify({
            "status": 500,
            "isError": True,
            "message": f"Error serving document: {str(e)}"
        }), 500

@employee_bp.get('/<int:employee_id>', responses={"200": EmployeeResponseSchema, "404": ErrorResponse}, security=[{"jwt": []}])
@jwt_required()
def get_employee(path: EmployeeIdPath):
    """Get employee by ID"""
    try:
        employee = Employee.query.get(path.employee_id)
        if not employee:
            return jsonify({
                "status": 404,
                "isError": True,
                "message": "Employee not found"
            }), 404
        
        employee_data = employee.to_dict()
        
        # Add company name if not already included
        if 'company_name' not in employee_data or not employee_data['company_name']:
            company = Company.query.get(employee.company_id)
            if company:
                employee_data['company_name'] = company.name
        
        return jsonify(employee_data), 200
        
    except Exception as e:
        return jsonify({
            "status": 500,
            "isError": True,
            "message": f"Error fetching employee: {str(e)}"
        }), 500

@employee_bp.put('/<int:employee_id>', responses={"200": EmployeeResponseSchema, "404": ErrorResponse, "500": ErrorResponse}, security=[{"jwt": []}])
@jwt_required()
def update_employee(path: EmployeeIdPath, body: EmployeeUpdateSchema):
    """Update employee details"""
    try:
        employee = Employee.query.get(path.employee_id)
        if not employee:
            return jsonify({
                "status": 404,
                "isError": True,
                "message": "Employee not found"
            }), 404

        # Update fields if provided
        update_data = body.model_dump(exclude_unset=True, exclude_none=True)
        
        # Remove force_update from update data as it's not an employee field
        update_data.pop('force_update', None)
        
        # Parse date fields if provided
        date_fields = ['date_of_birth', 'start_date', 'end_date', 'probation_end_date', 'contract_end_date', 'work_permit_valid_from', 'work_permit_valid_to']
        for field in date_fields:
            if field in update_data and update_data[field]:
                try:
                    update_data[field] = parse_date(update_data[field])
                except ValueError as e:
                    return jsonify({
                        "status": 400,
                        "isError": True,
                        "message": f"Invalid date format for {field}: {str(e)}"
                    }), 400

        # Handle identity type changes and validation based on nationality
        current_nationality = update_data.get('nationality', employee.nationality)
        is_zambian = current_nationality.lower() == 'zambian' if current_nationality else employee.nationality.lower() == 'zambian'
        
        if 'identity_type' in update_data:
            new_identity_type = update_data['identity_type']
            if is_zambian:
                # For Zambians, only NRC is allowed
                if new_identity_type == 'Work Permit':
                    return jsonify({
                        "status": 400,
                        "isError": True,
                        "message": "Zambian employees must use NRC as identity type"
                    }), 400
                
                # Clear work permit fields when switching to NRC
                update_data['work_permit_number'] = None
                update_data['work_permit_valid_from'] = None
                update_data['work_permit_valid_to'] = None
                update_data['work_permit_expiry_notified'] = False
                
                if not update_data.get('national_id') and not employee.national_id:
                    return jsonify({
                        "status": 400,
                        "isError": True,
                        "message": "National ID is required for Zambian employees"
                    }), 400
            else:
                # For non-Zambians, work permit is required
                if new_identity_type != 'Work Permit':
                    return jsonify({
                        "status": 400,
                        "isError": True,
                        "message": "Non-Zambian employees must use Work Permit as identity type"
                    }), 400
                
                # Clear national_id when switching to Work Permit
                update_data['national_id'] = None
                
                if not update_data.get('work_permit_number') and not employee.work_permit_number:
                    return jsonify({
                        "status": 400,
                        "isError": True,
                        "message": "Work Permit Number is required for non-Zambian employees"
                    }), 400
                if not update_data.get('work_permit_valid_from') and not employee.work_permit_valid_from:
                    return jsonify({
                        "status": 400,
                        "isError": True,
                        "message": "Work Permit valid from date is required for non-Zambian employees"
                    }), 400
                if not update_data.get('work_permit_valid_to') and not employee.work_permit_valid_to:
                    return jsonify({
                        "status": 400,
                        "isError": True,
                        "message": "Work Permit valid to date is required for non-Zambian employees"
                    }), 400

        # Update employee fields
        for key, value in update_data.items():
            if hasattr(employee, key):
                setattr(employee, key, value)

        employee.updated_at = datetime.now()
        db.session.commit()

        # Return updated employee
        employee_data = employee.to_dict()
        company = Company.query.get(employee.company_id)
        if company:
            employee_data['company_name'] = company.name

        return jsonify(employee_data), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({
            "status": 500,
            "isError": True,
            "message": f"Error updating employee: {str(e)}"
        }), 500

@employee_bp.delete('/<int:employee_id>', responses={"200": SuccessResponse, "404": ErrorResponse, "500": ErrorResponse}, security=[{"jwt": []}])
@jwt_required()
def delete_employee(path: EmployeeIdPath):
    """Delete employee"""
    try:
        employee = Employee.query.get(path.employee_id)
        if not employee:
            return jsonify({
                "status": 404,
                "isError": True,
                "message": "Employee not found"
            }), 404

        db.session.delete(employee)
        db.session.commit()

        return jsonify({
            "message": "Employee deleted successfully",
            "status": 200,
            "isError": False
        }), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({
            "status": 500,
            "isError": True,
            "message": f"Error deleting employee: {str(e)}"
        }), 500

@employee_bp.post('/<int:employee_id>/generate-documents', responses={"200": DocumentGenerationResponse, "404": ErrorResponse, "500": ErrorResponse}, security=[{"jwt": []}])
@jwt_required()
def generate_employee_documents(path: EmployeeIdPath):
    """Generate onboarding documents for an existing employee"""
    try:
        employee = Employee.query.get(path.employee_id)
        if not employee:
            return jsonify({
                "status": 404,
                "isError": True,
                "message": "Employee not found"
            }), 404

        company = Company.query.get(employee.company_id)
        if not company:
            return jsonify({
                "status": 404,
                "isError": True,
                "message": "Company not found"
            }), 404

        current_user_id = get_jwt_identity()
        
        # Generate documents
        generated_docs = generate_documents_from_templates(employee, company)
        
        # Save document records to database
        new_joiner_url = None
        contract_url = None
        
        if generated_docs.get('new_joiner_form'):
            drive_file_id = generated_docs.get('new_joiner_form_drive_id')
            new_joiner_doc = save_document_to_database(
                employee_id=employee.id,
                document_type='new_joiner_form',
                file_path=generated_docs['new_joiner_form'],
                uploaded_by=current_user_id,
                document_name=f"New Joiner Form - {employee.first_name} {employee.last_name}",
                comments="Employee onboarding form with personal and employment details",
                drive_file_id=drive_file_id
            )
            if new_joiner_doc:
                new_joiner_url = new_joiner_doc.file_url
        
        if generated_docs.get('employment_contract'):
            # Calculate contract expiry date
            expiry_date = None
            if employee.contract_end_date:
                expiry_date = employee.contract_end_date
            elif employee.start_date:
                # Default 1-year contract if no end date specified
                expiry_date = employee.start_date + timedelta(days=365)
            
            drive_file_id = generated_docs.get('employment_contract_drive_id')
            contract_doc = save_document_to_database(
                employee_id=employee.id,
                document_type='employment_contract',
                file_path=generated_docs['employment_contract'],
                uploaded_by=current_user_id,
                document_name=f"Employment Contract - {employee.first_name} {employee.last_name}",
                comments="Formal employment agreement with terms and conditions",
                expiry_date=expiry_date,
                drive_file_id=drive_file_id
            )
            if contract_doc:
                contract_url = contract_doc.file_url

        return jsonify({
            "employee_id": employee.id,
            "message": "Documents generated successfully",
            "new_joiner_form_url": new_joiner_url,
            "contract_url": contract_url
        }), 200

    except Exception as e:
        return jsonify({
            "status": 500,
            "isError": True,
            "message": f"Error generating documents: {str(e)}"
        }), 500

# ---------------------- BULK UPLOAD ROUTES ---------------------- #
@employee_bp.post(
    '/bulk-upload',
    responses={"200": BulkEmployeeResponseSchema, "400": ErrorResponse, "403": ErrorResponse, "500": ErrorResponse},
    security=[{"jwt": []}]
)
@jwt_required()
def bulk_create_employees(body: BulkEmployeeCreateSchema):
    """Bulk create multiple employees with Google Drive document handling"""
    try:
        current_user_id = get_jwt_identity()
        
        # Check HR admin access
        jwt_data = get_jwt()
        roles = jwt_data.get('roles', [])
        
        if 'hr_admin' not in roles and 'admin' not in roles:
            return jsonify({
                "status": 403,
                "isError": True,
                "message": "HR admin or Super Admin access required for bulk operations"
            }), 403

        if not body.employees:
            return jsonify({
                "status": 400,
                "isError": True,
                "message": "No employees provided for bulk upload"
            }), 400

        # Limit bulk upload size
        if len(body.employees) > 100:
            return jsonify({
                "status": 400,
                "isError": True,
                "message": "Bulk upload limited to 100 employees at a time"
            }), 400

        results = {
            "total_processed": len(body.employees),
            "successful": 0,
            "failed": 0,
            "errors": [],
            "successful_employees": []
        }

        # Process each employee
        for index, employee_data in enumerate(body.employees):
            try:
                # Validate company exists
                company = Company.query.get(employee_data.company_id)
                if not company:
                    error_msg = f"Company with ID {employee_data.company_id} not found"
                    if not body.skip_errors:
                        raise ValueError(error_msg)
                    results["errors"].append({
                        "index": index,
                        "employee_name": f"{employee_data.first_name} {employee_data.last_name}",
                        "error": error_msg
                    })
                    results["failed"] += 1
                    continue

                # Validate supervisor exists if provided
                if employee_data.supervisor_id:
                    supervisor = Employee.query.get(employee_data.supervisor_id)
                    if not supervisor:
                        error_msg = f"Supervisor with ID {employee_data.supervisor_id} not found"
                        if not body.skip_errors:
                            raise ValueError(error_msg)
                        results["errors"].append({
                            "index": index,
                            "employee_name": f"{employee_data.first_name} {employee_data.last_name}",
                            "error": error_msg
                        })
                        results["failed"] += 1
                        continue

                # Generate employee ID
                employee_id = generate_employee_id(employee_data.company_id)

                # Check if generated employee ID already exists
                existing_employee = Employee.query.filter_by(employee_id=employee_id).first()
                if existing_employee:
                    employee_id = f"EMP{int(datetime.now().timestamp())}_{index}"

                # Check if email already exists
                if employee_data.email:
                    existing_employee = Employee.query.filter_by(email=employee_data.email).first()
                    if existing_employee:
                        error_msg = f"An employee with email {employee_data.email} already exists"
                        if not body.skip_errors:
                            raise ValueError(error_msg)
                        results["errors"].append({
                            "index": index,
                            "employee_name": f"{employee_data.first_name} {employee_data.last_name}",
                            "error": error_msg
                        })
                        results["failed"] += 1
                        continue

                # Check identity document uniqueness based on nationality
                is_zambian = employee_data.nationality.lower() == 'zambian'
                
                if is_zambian:
                    # For Zambians, check national_id uniqueness
                    if employee_data.national_id:
                        existing_national_id = Employee.query.filter_by(national_id=employee_data.national_id).first()
                        if existing_national_id:
                            error_msg = f"An employee with National ID {employee_data.national_id} already exists"
                            if not body.skip_errors:
                                raise ValueError(error_msg)
                            results["errors"].append({
                                "index": index,
                                "employee_name": f"{employee_data.first_name} {employee_data.last_name}",
                                "error": error_msg
                            })
                            results["failed"] += 1
                            continue
                else:
                    # For non-Zambians, check work permit uniqueness
                    if employee_data.work_permit_number:
                        existing_work_permit = Employee.query.filter_by(work_permit_number=employee_data.work_permit_number).first()
                        if existing_work_permit:
                            error_msg = f"An employee with Work Permit Number {employee_data.work_permit_number} already exists"
                            if not body.skip_errors:
                                raise ValueError(error_msg)
                            results["errors"].append({
                                "index": index,
                                "employee_name": f"{employee_data.first_name} {employee_data.last_name}",
                                "error": error_msg
                            })
                            results["failed"] += 1
                            continue

                # Parse dates using the updated parse_date function
                date_of_birth = parse_date(employee_data.date_of_birth)
                start_date = parse_date(employee_data.start_date)
                end_date = parse_date(employee_data.end_date) if employee_data.end_date else None
                probation_end_date = parse_date(employee_data.probation_end_date) if employee_data.probation_end_date else None
                contract_end_date = parse_date(employee_data.contract_end_date) if employee_data.contract_end_date else None
                
                # Parse work permit dates for non-Zambians
                work_permit_valid_from = None
                work_permit_valid_to = None
                if not is_zambian and employee_data.identity_type == 'Work Permit':
                    work_permit_valid_from = parse_date(employee_data.work_permit_valid_from) if employee_data.work_permit_valid_from else None
                    work_permit_valid_to = parse_date(employee_data.work_permit_valid_to) if employee_data.work_permit_valid_to else None

                # Handle null salary
                salary = employee_data.salary if employee_data.salary is not None else 0.0

                # Create employee
                employee = Employee(
                    employee_id=employee_id,
                    first_name=employee_data.first_name,
                    last_name=employee_data.last_name,
                    email=employee_data.email,
                    phone=employee_data.phone or "Not Provided",
                    personal_email=employee_data.personal_email,
                    date_of_birth=date_of_birth,
                    
                    # Identity fields
                    national_id=employee_data.national_id,
                    work_permit_number=employee_data.work_permit_number,
                    identity_type=employee_data.identity_type,
                    work_permit_valid_from=work_permit_valid_from,
                    work_permit_valid_to=work_permit_valid_to,
                    work_permit_expiry_notified=False,
                    nationality=employee_data.nationality,
                    
                    # Additional fields
                    middle_name=employee_data.middle_name,
                    napsa_number=employee_data.napsa_number,
                    nhima_number=employee_data.nhima_number,
                    tpin=employee_data.tax_id,
                    account_number=employee_data.bank_account,
                    sort_code=employee_data.sort_code,
                    next_of_kin=employee_data.next_of_kin,
                    physical_address=employee_data.physical_address,
                    
                    gender=employee_data.gender,
                    marital_status=employee_data.marital_status,
                    address=employee_data.address,
                    emergency_contact_name=employee_data.emergency_contact_name or "Not Provided",
                    emergency_contact_phone=employee_data.emergency_contact_phone or "Not Provided",
                    emergency_contact_relationship=employee_data.emergency_contact_relationship or "Not Provided",
                    company_id=employee_data.company_id,
                    department=employee_data.department,
                    position=employee_data.position,
                    employment_type=employee_data.employment_type,
                    employment_status=employee_data.employment_status,
                    start_date=start_date,
                    end_date=end_date,
                    probation_end_date=probation_end_date,
                    contract_end_date=contract_end_date,
                    supervisor_id=employee_data.supervisor_id,
                    work_location=employee_data.work_location,
                    salary=salary,
                    salary_currency=employee_data.salary_currency,
                    payment_frequency=employee_data.payment_frequency,
                    bank_name=employee_data.bank_name,
                    bank_account=employee_data.bank_account,
                    tax_id=employee_data.tax_id,
                    pension_number=employee_data.pension_number,
                    created_by=int(current_user_id)
                )
                
                db.session.add(employee)
                db.session.flush()
                
                # Handle uploaded documents if provided with Google Drive
                if employee_data.documents:
                    try:
                        saved_documents = handle_employee_documents(
                            employee_id=employee.id,
                            documents_data=employee_data.documents,
                            uploaded_by=current_user_id,
                            employee=employee,
                            company=company
                        )
                        print(f"DEBUG: Saved {len(saved_documents)} documents for employee {employee.id}")
                    except Exception as doc_error:
                        print(f"Warning: Document processing failed but employee created: {str(doc_error)}")
                        # Continue with employee creation even if documents fail
                
                # Generate template documents if requested
                if employee_data.generate_documents:
                    try:
                        documents_data = {
                            "nhima_number": employee_data.nhima_number,
                            "bank_branch": employee_data.bank_branch,
                            "sort_code": employee_data.sort_code,
                            "spouse_name": employee_data.spouse_name,
                            "children": employee_data.children,
                            "housing_allowance": employee_data.housing_allowance or 0,
                            "transport_allowance": employee_data.transport_allowance or 0,
                            "lunch_allowance": employee_data.lunch_allowance or 0,
                            "fuel_allowance": employee_data.fuel_allowance or 0,
                            "phone_allowance": employee_data.phone_allowance or 0,
                            "company_vehicle": employee_data.company_vehicle,
                            "company_phone": employee_data.company_phone,
                            "company_accommodation": employee_data.company_accommodation,
                            "working_hours": employee_data.working_hours
                        }
                        
                        generated_docs = generate_documents_from_templates(employee, company, documents_data)
                        
                        # Save documents to database with Google Drive support
                        if generated_docs.get('new_joiner_form'):
                            drive_file_id = generated_docs.get('new_joiner_form_drive_id')
                            save_document_to_database(
                                employee_id=employee.id,
                                document_type='new_joiner_form',
                                file_path=generated_docs['new_joiner_form'],
                                uploaded_by=current_user_id,
                                document_name=f"New Joiner Form - {employee.first_name} {employee.last_name}",
                                comments="Employee onboarding form with personal and employment details",
                                drive_file_id=drive_file_id
                            )
                        
                        if generated_docs.get('employment_contract'):
                            expiry_date = None
                            if employee.contract_end_date:
                                expiry_date = employee.contract_end_date
                            elif employee.start_date:
                                expiry_date = employee.start_date + timedelta(days=365)
                            
                            drive_file_id = generated_docs.get('employment_contract_drive_id')
                            save_document_to_database(
                                employee_id=employee.id,
                                document_type='employment_contract',
                                file_path=generated_docs['employment_contract'],
                                uploaded_by=current_user_id,
                                document_name=f"Employment Contract - {employee.first_name} {employee.last_name}",
                                comments="Formal employment agreement with terms and conditions",
                                expiry_date=expiry_date,
                                drive_file_id=drive_file_id
                            )
                            
                    except Exception as doc_error:
                        print(f"Template document generation failed but employee created: {str(doc_error)}")
                
                # Update company employee count
                company.employee_count = Employee.query.filter_by(
                    company_id=employee_data.company_id
                ).filter(
                    Employee.employment_status.in_(['Active', 'Probation'])
                ).count()
                
                # Create audit log
                audit = AuditLog(
                    employee_id=employee.id,
                    action="CREATE",
                    performed_by=current_user_id,
                    details=f"Employee {employee.employee_id} created via bulk upload with nationality: {employee.nationality}, identity type: {employee.identity_type}. Documents stored in Google Drive."
                )
                db.session.add(audit)
                
                # Add to successful results
                employee_dict = employee.to_dict()
                if 'company_name' not in employee_dict or not employee_dict['company_name']:
                    employee_dict['company_name'] = company.name
                
                results["successful_employees"].append({
                    "index": index,
                    "employee_id": employee.employee_id,
                    "employee_db_id": employee.id,
                    "name": f"{employee.first_name} {employee.last_name}",
                    "email": employee.email,
                    "position": employee.position,
                    "department": employee.department,
                    "nationality": employee.nationality
                })
                results["successful"] += 1
                
                print(f"Successfully created employee {index + 1}/{len(body.employees)}: {employee.employee_id}")

            except Exception as e:
                db.session.rollback()
                error_msg = f"Error creating employee {employee_data.first_name} {employee_data.last_name}: {str(e)}"
                print(f"Bulk upload error at index {index}: {error_msg}")
                
                results["errors"].append({
                    "index": index,
                    "employee_name": f"{employee_data.first_name} {employee_data.last_name}",
                    "error": error_msg
                })
                results["failed"] += 1
                
                if not body.skip_errors:
                    # If not skipping errors, stop processing
                    return jsonify({
                        "status": 500,
                        "isError": True,
                        "message": f"Bulk upload failed at employee {index + 1}: {error_msg}",
                        "partial_results": results
                    }), 500
        
        # Commit all successful creations
        if results["successful"] > 0:
            db.session.commit()
        
        # Final response
        message = f"Bulk upload completed: {results['successful']} successful, {results['failed']} failed"
        if results["failed"] > 0 and body.skip_errors:
            message += " (errors were skipped)"
        
        response_data = {
            "message": message,
            "status": 200,
            "isError": False,
            **results
        }
        
        return jsonify(response_data), 200

    except Exception as e:
        db.session.rollback()
        print(f"Error in bulk_create_employees: {str(e)}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        return jsonify({
            "status": 500,
            "isError": True,
            "message": f"Error in bulk upload: {str(e)}"
        }), 500

# Also add this helper endpoint for bulk upload template
@employee_bp.get('/bulk-upload-template', responses={"200": None}, security=[{"jwt": []}])
@jwt_required()
def get_bulk_upload_template():
    """Get template for bulk employee upload"""
    template = {
        "employees": [
            {
                "first_name": "Moses",
                "middle_name": "K.",
                "last_name": "Jasi",
                "email": "moses.jasi@example.com",
                "phone": "+260971234567",
                "personal_email": "mjasi.personal@example.com",
                "gender": "Male",
                "marital_status": "Single",
                "date_of_birth": "1992-06-15",
                "address": "Plot 12, Lusaka",
                "physical_address": "Lusaka, Zambia",
                
                "identity_type": "NRC",
                "national_id": "123456/12/1",
                "work_permit_number": None,
                "work_permit_valid_from": None,
                "work_permit_valid_to": None,
                "nationality": "Zambian",
                
                "company_id": 1,
                "department": "IT Department",
                "position": "Software Engineer",
                "employment_type": "Fixed Term",
                "employment_status": "Active",
                "start_date": "2025-11-01",
                "end_date": None,
                "probation_end_date": None,
                "contract_end_date": None,
                "supervisor_id": None,
                "work_location": "Head Office",
                "salary": 12000.0,
                "salary_currency": "ZMW",
                "payment_frequency": "Monthly",
                
                "bank_name": "Stanbic Bank",
                "bank_branch": "Cairo Road",
                "bank_account": "1234567890",
                "sort_code": "090009",
                "account_number": "1234567890",
                
                "tax_id": "TPIN123456",
                "pension_number": "NAPSA98765",
                "nhima_number": "NHIMA55555",
                "napsa_number": "NAPSA98765",
                
                "next_of_kin": "Emmanuel Chanda",
                "emergency_contact_name": "Emmanuel Chanda",
                "emergency_contact_relationship": "Brother",
                "emergency_contact_phone": "+260977000999",
                
                "generate_documents": True,
                
                "documents": {
                    "profilePhoto": "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgA...",  # Optional
                    "nrcCopy": "data:application/pdf;base64,JVBERi0xLjQKJeLjz9M...",  # Optional
                    "cv": "data:application/pdf;base64,JVBERi0xLjQKJeLjz9M...",  # Optional
                    "offerLetter": "data:application/pdf;base64,JVBERi0xLjQKJeLjz9M...",  # Optional
                    "certificates": [  # Optional
                        "data:application/pdf;base64,JVBERi0xLjQKJeLjz9M...",
                        "data:image/jpeg;base64,/9j/4AAQSkZJRgABAQAAAQABAAD..."
                    ]
                }
            }
        ],
        "skip_errors": False,
        "send_notifications": False
    }
    
    return jsonify({
        "template": template,
        "instructions": {
            "required_fields": ["first_name", "last_name", "date_of_birth", "gender", "company_id", "position", "department", "employment_type", "start_date", "nationality"],
            "optional_fields": "All other fields are optional including documents",
            "identity_documents": "Provide national_id for Zambians or work_permit_number with dates for non-Zambians",
            "employment_types": "Allowed values: Full-time, Part-time, Contract, Permanent, Fixed Term, Intern, Apprentice, Consultant",
            "nationality": "Use 'Zambian' for Zambian citizens, any other nationality for non-Zambians",
            "documents": "All document fields are optional. Supported formats: PNG, JPG, JPEG, PDF",
            "max_file_size": "10MB per document",
            "max_employees": "100 per bulk upload"
        }
    }), 200

# ---------------------- WORK PERMIT EXPIRY NOTIFICATION ROUTES ---------------------- #
@employee_bp.get('/work-permit-expiry-notifications', responses={"200": WorkPermitExpiryResponse, "500": ErrorResponse}, security=[{"jwt": []}])
@jwt_required()
def get_expiring_work_permits():
    """Get employees with expiring or expired work permits"""
    try:
        today = datetime.now().date()
        thirty_days_from_now = today + timedelta(days=30)
        
        # Employees with expired work permits (non-Zambians only)
        expired_employees = Employee.query.filter(
            Employee.identity_type == 'Work Permit',
            Employee.nationality != 'Zambian',
            Employee.work_permit_valid_to < today,
            Employee.work_permit_expiry_notified == False,
            Employee.employment_status.in_(['Active', 'Probation'])
        ).all()
        
        # Employees with work permits expiring within 30 days (non-Zambians only)
        expiring_soon_employees = Employee.query.filter(
            Employee.identity_type == 'Work Permit',
            Employee.nationality != 'Zambian',
            Employee.work_permit_valid_to >= today,
            Employee.work_permit_valid_to <= thirty_days_from_now,
            Employee.work_permit_expiry_notified == False,
            Employee.employment_status.in_(['Active', 'Probation'])
        ).all()
        
        expired_data = [emp.to_dict() for emp in expired_employees]
        expiring_soon_data = [emp.to_dict() for emp in expiring_soon_employees]
        
        return jsonify({
            "expired_work_permits": expired_data,
            "expiring_soon_work_permits": expiring_soon_data,
            "expired_count": len(expired_data),
            "expiring_soon_count": len(expiring_soon_data),
            "check_date": today.isoformat()
        }), 200
        
    except Exception as e:
        return jsonify({
            "status": 500,
            "isError": True,
            "message": f"Error checking expiring work permits: {str(e)}"
        }), 500

@employee_bp.post('/<int:employee_id>/mark-work-permit-notified', responses={"200": SuccessResponse, "404": ErrorResponse, "500": ErrorResponse}, security=[{"jwt": []}])
@jwt_required()
def mark_work_permit_expiry_notified(path: EmployeeIdPath):
    """Mark work permit expiry as notified"""
    try:
        employee = Employee.query.get(path.employee_id)
        if not employee:
            return jsonify({
                "status": 404,
                "isError": True,
                "message": "Employee not found"
            }), 404

        if employee.identity_type != 'Work Permit' or employee.nationality.lower() == 'zambian':
            return jsonify({
                "status": 400,
                "isError": True,
                "message": "Employee does not have a work permit or is Zambian"
            }), 400

        employee.work_permit_expiry_notified = True
        employee.updated_at = datetime.now()
        db.session.commit()

        return jsonify({
            "message": "Work permit expiry notification marked as sent",
            "status": 200,
            "isError": False
        }), 200

    except Exception as e:
        db.session.rollback()
        return jsonify({
            "status": 500,
            "isError": True,
            "message": f"Error updating work permit notification status: {str(e)}"
        }), 500